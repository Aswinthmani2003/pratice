import streamlit as st
from docx import Document
from datetime import datetime
import os
import platform
import subprocess

# Function to replace placeholders in Word document
def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document, including paragraphs and tables."""
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a single paragraph, handling split runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""  # Clear all runs
            paragraph.runs[0].text = full_text  # Add the replaced text back
    
    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        replace_in_paragraph(para, key, value)
    return doc

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts DOCX to PDF while retaining formatting."""
    if platform.system() == "Windows":
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
    elif platform.system() == "Linux":
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_path])
    return pdf_path

# Streamlit UI
st.title("Contract Agreement Generator")

# Input fields for placeholders
party1_name = st.text_input("Party 1 Name")
party1_address = st.text_input("Party 1 Address")
party1_email = st.text_input("Party 1 Email")
party1_phone = st.text_input("Party 1 Phone")

placeholders = {
    "<<Name>>": party1_name,
    "<<Address>>": party1_address,
    "<<Email>>": party1_email,
    "<<PhoneNumber>>": party1_phone,
}

# Generate contract
if st.button("Generate Contract Agreement"):
    try:
        doc_path = "sample_contract.docx"
        doc = Document(doc_path)
        doc = replace_placeholders(doc, placeholders)
        
        # Save as DOCX
        docx_output_path = "Generated_Contract.docx"
        doc.save(docx_output_path)

        # Convert to PDF
        pdf_output_path = "Generated_Contract.pdf"
        convert_docx_to_pdf(docx_output_path, pdf_output_path)
        
        st.success("Contract Agreement generated successfully!")
        
        col1, col2 = st.columns(2)
        
        with col1:
            with open(docx_output_path, "rb") as file:
                st.download_button("Download as Word", file, file_name="Contract_Agreement.docx")
        
        with col2:
            with open(pdf_output_path, "rb") as file:
                st.download_button("Download as PDF", file, file_name="Contract_Agreement.pdf")
    
    except Exception as e:
        st.error(f"Error: {e}")
